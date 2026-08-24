"""Word 解析：抽取段落与表格文本。"""
from pathlib import Path

from docx import Document


def parse(path: Path) -> str:
    """把 docx 的段落和表格转成纯文本。"""
    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(", ".join(cells))
    return "\n".join(lines)
